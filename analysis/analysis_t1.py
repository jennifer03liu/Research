import pandas as pd
import semopy
import numpy as np
import os
import glob
from scipy.stats import pearsonr
from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH

# ==========================================
# 0. Helper Functions
# ==========================================

def calculate_cronbach_alpha(df):
    """
    Calculates Cronbach's Alpha for a given DataFrame of items.
    """
    df_corr = df.corr()
    N = df.shape[1]
    
    rs = np.array([])
    for i, col1 in enumerate(df_corr.columns):
        for j, col2 in enumerate(df_corr.columns):
            if i > j:
                rs = np.append(rs, df_corr.iloc[i, j])
    
    mean_r = np.mean(rs)
    alpha = (N * mean_r) / (1 + (N - 1) * mean_r)
    return alpha

def get_significance_stars(p_value):
    if p_value < 0.001: return '***'
    elif p_value < 0.01: return '**'
    elif p_value < 0.05: return '*'
    else: return ''

def save_to_word(stats_df, corr_df, cfa_loadings, cfa_fit, filename="Research_Result.docx"):
    print(f"\nGenerating Word Report: {filename}...")
    doc = Document()
    
    # Title
    head = doc.add_heading('Research Analysis Report', 0)
    head.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    def add_table_to_doc(df, title):
        doc.add_heading(title, level=2)
        table = doc.add_table(rows=1, cols=len(df.columns))
        table.style = 'Table Grid' # Basic style, can be 'Table Normal' for no borders
        
        # Header
        hdr_cells = table.rows[0].cells
        for i, col_name in enumerate(df.columns):
            hdr_cells[i].text = str(col_name)
            
        # Body
        for index, row in df.iterrows():
            row_cells = table.add_row().cells
            for i, val in enumerate(row):
                row_cells[i].text = str(val)
        
        doc.add_paragraph("\n")

    # 1. Descriptive & Reliability
    add_table_to_doc(stats_df, "Table 1. Descriptive Statistics and Reliability")

    # 2. Correlation Matrix
    add_table_to_doc(corr_df, "Table 2. Correlation Matrix")

    # 3. Model Fit
    # Select key columns
    fit_cols = ['RMSEA', 'CFI', 'TLI', 'chi2', 'chi2 p-value', 'DoF']
    # Filter columns that exist in cfa_fit
    existing_cols = [c for c in fit_cols if c in cfa_fit.columns]
    fit_summary = cfa_fit[existing_cols]
    add_table_to_doc(fit_summary, "Table 3. Model Fit Indices")

    # 4. Factor Loadings
    # Filter for loadings (op is '=~')
    loadings = cfa_loadings[cfa_loadings['op'] == '=~'].copy()
    # Select columns
    loadings = loadings[['lval', 'rval', 'Estimate', 'p-value', 'Std. Err']]
    # Round
    loadings['Estimate'] = loadings['Estimate'].apply(lambda x: round(x, 3))
    loadings['p-value'] = loadings['p-value'].apply(lambda x: round(x, 3))
    loadings['Std. Err'] = loadings['Std. Err'].apply(lambda x: round(x, 3) if x != '-' else '-')
    
    add_table_to_doc(loadings, "Table 4. Factor Loadings")

    doc.save(filename)
    print(f"Report saved to {filename}")

# ==========================================
# 1. Load Data
# ==========================================
# Define Scale Items
scale_items = {
    "HCP": ["HCP1", "HCP2", "HCP3", "HCP4_R", "HCP5", "HCP6_R"],
    "JCP": ["JCP1_R", "JCP2_R", "JCP3_R", "JCP4_R", "JCP5_R", "JCP6"],
    "PP":  ["PP1", "PP2", "PP3", "PP4", "PP5", "PP6"],
    "DP":  ["DP1", "DP2", "DP3", "DP4", "DP5"],
    "CI":  ["CI1", "CI2", "CI3", "CI4", "CI5", "CI6", "CI7", "CI8"]
}

# Find latest CSV file matching pattern
search_pattern = 'T1_*_SPSS.csv'
list_of_files = glob.glob(search_pattern)

if list_of_files:
    # Use os.path.getmtime to sort by modification time, newest last
    latest_file = max(list_of_files, key=os.path.getmtime)
    print(f"Found {len(list_of_files)} data files. Using the latest: {latest_file}")
    data = pd.read_csv(latest_file)
else:
    print(f"Warning: No files matching '{search_pattern}' found in current directory.")
    print("Generating MOCK data for demonstration...")
    np.random.seed(42)
    N = 377
    data = pd.DataFrame()
    for scale, items in scale_items.items():
        # Mocking data
        f = np.random.normal(0, 1, N)
        for i, col_name in enumerate(items):
            item_val = 0.7 * f + np.random.normal(0, 0.5, N)
            item_val = (item_val * 1.0) + 3.5
            item_val = np.clip(item_val, 1, 6)
            data[col_name] = item_val
    print("Generated Mock Columns:", data.columns.tolist())

# ==========================================
# 2. Reliability & Descriptive Statistics
# ==========================================
print("\n--- Reliability & Descriptive Statistics ---")
stats_list = []
scale_scores = pd.DataFrame()

for scale, items in scale_items.items():
    missing_cols = [c for c in items if c not in data.columns]
    if missing_cols: continue
        
    df_subset = data[items]
    alpha = calculate_cronbach_alpha(df_subset)
    scale_score = df_subset.mean(axis=1)
    scale_scores[scale] = scale_score
    
    stats_list.append({
        "Variable": scale,
        "Items": len(items),
        "Mean": round(scale_score.mean(), 2),
        "SD": round(scale_score.std(), 2),
        "Alpha": round(alpha, 2)
    })

stats_df = pd.DataFrame(stats_list)
print(stats_df)
stats_df.to_csv("descriptive_reliability.csv", index=False)

# ==========================================
# 3. Correlation Matrix with Stars
# ==========================================
print("\n--- Correlation Matrix ---")
cols = stats_df['Variable'].tolist()
corr_data = []

for r in cols:
    row_data = {'Variable': r}
    for c in cols:
        if r == c:
            row_data[c] = '1.00'
        else:
            r_val, p_val = pearsonr(scale_scores[r], scale_scores[c])
            star = get_significance_stars(p_val)
            row_data[c] = f"{r_val:.2f}{star}"
    corr_data.append(row_data)

corr_df = pd.DataFrame(corr_data)
print(corr_df)
corr_df.to_csv("correlation_matrix.csv", index=False)

# ==========================================
# 4. Run Analysis (CFA) & Export Word
# ==========================================
model_desc = """
    HCP =~ HCP1 + HCP2 + HCP3 + HCP4_R + HCP5 + HCP6_R
    JCP =~ JCP1_R + JCP2_R + JCP3_R + JCP4_R + JCP5_R
    PP =~ PP1 + PP2 + PP3 + PP4 + PP5 + PP6
    DP =~ DP1 + DP2 + DP3 + DP4 + DP5
    CI =~ CI1 + CI2 + CI3 + CI4 + CI5 + CI6 + CI7 + CI8
"""

print("\nRunning CFA Model...")
cfa_loadings = pd.DataFrame()
cfa_fit = pd.DataFrame()

try:
    model = semopy.Model(model_desc)
    model.fit(data)
    
    cfa_loadings = model.inspect()
    cfa_fit = semopy.calc_stats(model).T
    
    print("\n--- Model Fit Indices ---")
    print(cfa_fit) 
    
    cfa_loadings.to_csv("cfa_results_loadings.csv")
    cfa_fit.to_csv("cfa_results_fit.csv")

except Exception as e:
    print(f"\nError running model: {e}")

# EXPORT TO WORD
save_to_word(stats_df, corr_df, cfa_loadings, cfa_fit)
